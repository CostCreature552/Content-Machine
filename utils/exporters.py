"""
Export utilities for blog content.
Supports Markdown (.md), PDF (.pdf), and DOCX (.docx) formats.
"""

import io
import re
from typing import Optional


def _sanitize_for_pdf(text: str) -> str:
    """
    Replace Unicode characters that are not supported by fpdf2's
    built-in Helvetica font with ASCII equivalents.
    """
    replacements = {
        "\u2018": "'",   # left single quote
        "\u2019": "'",   # right single quote (curly apostrophe)
        "\u201C": '"',   # left double quote
        "\u201D": '"',   # right double quote
        "\u2013": "-",   # en dash
        "\u2014": "--",  # em dash
        "\u2026": "...", # ellipsis
        "\u00A0": " ",   # non-breaking space
        "\u200B": "",    # zero-width space
        "\u00AB": '"',   # left-pointing double angle quote
        "\u00BB": '"',   # right-pointing double angle quote
        "\u2022": "-",   # bullet (we re-add our own)
        "\u2032": "'",   # prime
        "\u2033": '"',   # double prime
        "\u00B7": "-",   # middle dot
        "\u2010": "-",   # hyphen
        "\u2011": "-",   # non-breaking hyphen
        "\u2012": "-",   # figure dash
        "\u00BD": "1/2", # vulgar fraction one half
        "\u00BC": "1/4", # vulgar fraction one quarter
        "\u00BE": "3/4", # vulgar fraction three quarters
        "\u00E9": "e",   # e acute
        "\u00E8": "e",   # e grave
        "\u00F1": "n",   # n tilde
        "\u00FC": "u",   # u umlaut
    }
    for unicode_char, ascii_char in replacements.items():
        text = text.replace(unicode_char, ascii_char)
    # Fallback: remove any remaining non-latin1 characters
    text = text.encode("latin-1", errors="replace").decode("latin-1")
    return text


def export_markdown(content: str, filename: str = "blog_article.md") -> bytes:
    """
    Export content as a Markdown file.

    Args:
        content: The Markdown content to export.
        filename: The filename (used for reference, not applied here).

    Returns:
        bytes: The content encoded as UTF-8 bytes.
    """
    return content.encode("utf-8")


def export_pdf(content: str, title: str = "Blog Article") -> bytes:
    """
    Export content as a PDF file.

    Args:
        content: The Markdown content to convert to PDF.
        title: The document title.

    Returns:
        bytes: The PDF file as bytes.
    """
    try:
        from fpdf import FPDF
    except ImportError:
        raise ImportError(
            "fpdf2 is required for PDF export. Install it with: pip install fpdf2"
        )

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Sanitize ALL content upfront to prevent Unicode encoding errors
    content = _sanitize_for_pdf(content)
    title = _sanitize_for_pdf(title)

    # Title
    pdf.set_font("Helvetica", "B", 20)
    pdf.set_text_color(33, 37, 41)
    pdf.cell(0, 15, title, new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(5)

    # Process content line by line
    lines = content.split("\n")

    for line in lines:
        stripped = line.strip()

        if not stripped:
            pdf.ln(3)
            continue

        # H1 heading
        if stripped.startswith("# ") and not stripped.startswith("## "):
            heading_text = stripped[2:].strip()
            pdf.set_font("Helvetica", "B", 18)
            pdf.set_text_color(33, 37, 41)
            pdf.ln(5)
            pdf.multi_cell(0, 10, heading_text)
            pdf.ln(3)

        # H2 heading
        elif stripped.startswith("## ") and not stripped.startswith("### "):
            heading_text = stripped[3:].strip()
            pdf.set_font("Helvetica", "B", 14)
            pdf.set_text_color(52, 58, 64)
            pdf.ln(4)
            pdf.multi_cell(0, 8, heading_text)
            pdf.ln(2)

        # H3 heading
        elif stripped.startswith("### "):
            heading_text = stripped[4:].strip()
            pdf.set_font("Helvetica", "B", 12)
            pdf.set_text_color(73, 80, 87)
            pdf.ln(3)
            pdf.multi_cell(0, 7, heading_text)
            pdf.ln(2)

        # Bullet points
        elif stripped.startswith("- ") or stripped.startswith("* "):
            bullet_text = stripped[2:].strip()
            # Remove markdown bold/italic
            bullet_text = re.sub(r"\*{1,3}(.*?)\*{1,3}", r"\1", bullet_text)
            pdf.set_font("Helvetica", "", 11)
            pdf.set_text_color(33, 37, 41)
            pdf.cell(5)
            pdf.multi_cell(0, 6, f"  -  {bullet_text}")
            pdf.ln(1)

        # Numbered list
        elif re.match(r"^\d+\.\s", stripped):
            list_text = re.sub(r"\*{1,3}(.*?)\*{1,3}", r"\1", stripped)
            pdf.set_font("Helvetica", "", 11)
            pdf.set_text_color(33, 37, 41)
            pdf.cell(5)
            pdf.multi_cell(0, 6, f"  {list_text}")
            pdf.ln(1)

        # Horizontal rule
        elif stripped in ("---", "***", "___"):
            pdf.ln(3)
            pdf.set_draw_color(200, 200, 200)
            pdf.line(10, pdf.get_y(), 200, pdf.get_y())
            pdf.ln(3)

        # Regular paragraph text
        else:
            # Remove markdown formatting
            clean_text = re.sub(r"\*{1,3}(.*?)\*{1,3}", r"\1", stripped)
            clean_text = re.sub(r"_{1,3}(.*?)_{1,3}", r"\1", clean_text)
            clean_text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", clean_text)
            clean_text = re.sub(r"`([^`]+)`", r"\1", clean_text)
            pdf.set_font("Helvetica", "", 11)
            pdf.set_text_color(33, 37, 41)
            pdf.multi_cell(0, 6, clean_text)
            pdf.ln(2)

    # Return PDF as bytes
    pdf_bytes = pdf.output()
    if isinstance(pdf_bytes, str):
        return pdf_bytes.encode("latin-1")
    return bytes(pdf_bytes)


def export_docx(content: str, title: str = "Blog Article") -> bytes:
    """
    Export content as a DOCX file.

    Args:
        content: The Markdown content to convert to DOCX.
        title: The document title.

    Returns:
        bytes: The DOCX file as bytes.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX export. Install it with: pip install python-docx"
        )

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(33, 37, 41)

    # Process content line by line
    lines = content.split("\n")

    for line in lines:
        stripped = line.strip()

        if not stripped:
            continue

        # H1 heading
        if stripped.startswith("# ") and not stripped.startswith("## "):
            heading_text = stripped[2:].strip()
            heading = doc.add_heading(heading_text, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # H2 heading
        elif stripped.startswith("## ") and not stripped.startswith("### "):
            heading_text = stripped[3:].strip()
            doc.add_heading(heading_text, level=2)

        # H3 heading
        elif stripped.startswith("### "):
            heading_text = stripped[4:].strip()
            doc.add_heading(heading_text, level=3)

        # Bullet points
        elif stripped.startswith("- ") or stripped.startswith("* "):
            bullet_text = stripped[2:].strip()
            bullet_text = re.sub(r"\*{1,3}(.*?)\*{1,3}", r"\1", bullet_text)
            doc.add_paragraph(bullet_text, style="List Bullet")

        # Numbered list
        elif re.match(r"^\d+\.\s", stripped):
            list_text = re.sub(r"^\d+\.\s", "", stripped)
            list_text = re.sub(r"\*{1,3}(.*?)\*{1,3}", r"\1", list_text)
            doc.add_paragraph(list_text, style="List Number")

        # Horizontal rule
        elif stripped in ("---", "***", "___"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("─" * 50)
            run.font.color.rgb = RGBColor(200, 200, 200)

        # Regular paragraph
        else:
            # Clean markdown formatting
            clean_text = re.sub(r"\*{1,3}(.*?)\*{1,3}", r"\1", stripped)
            clean_text = re.sub(r"_{1,3}(.*?)_{1,3}", r"\1", clean_text)
            clean_text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", clean_text)
            clean_text = re.sub(r"`([^`]+)`", r"\1", clean_text)
            doc.add_paragraph(clean_text)

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
